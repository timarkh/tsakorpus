import re
import math
from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.oxml.shared import OxmlElement, qn
from docx.enum.style import WD_STYLE_TYPE


class DocxExampleProcessor:
    """
    Contains methods for rendering search hit(s) as table(s) in a Word document.
    """
    rxStemGloss = re.compile('[ ,;:()]+')
    rxPuncR = re.compile('^[.,?!:;)"/\\-\\]”]+$')
    rxPuncL = re.compile('^(?:[/?!. ]*\\[|[#*(\\[“]+)$')
    rxLetters = re.compile('\\w')
    rxLeadingNewline = re.compile('^(\n|\\\\n) *', flags=re.DOTALL)
    rxTrailingNewline = re.compile(' *(\n|\\\\n)$', flags=re.DOTALL)
    rxGlossesNonGlosses = re.compile('([^$]+)')

    def __init__(self, settings):
        self.settings = settings

    @staticmethod
    def set_cell_margins(table, left=0, right=0):
        tc = table._element
        tblPr = tc.tblPr
        tblCellMar = OxmlElement('w:tblCellMar')
        kwargs = {"left": left, "right": right}
        for m in ["left", "right"]:
            node = OxmlElement("w:{}".format(m))
            node.set(qn('w:w'), str(kwargs.get(m)))
            node.set(qn('w:type'), 'dxa')
            tblCellMar.append(node)

        tblPr.append(tblCellMar)

    @staticmethod
    def p_no_margins(wordDoc, p, style='Normal'):
        p.style = wordDoc.styles[style]
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Cm(0)
        p.paragraph_format.space_after = Cm(0)

    def smallcaps_glosses(self, p, text, langProps):
        if 'gloss_regex' not in langProps:
            p.text = text.strip()
            return
        text = langProps['gloss_regex'].sub(lambda m: '$' + m.group(1) + '$', text)
        for run in self.rxGlossesNonGlosses.findall(text):
            if langProps['gloss_regex'].search(run) is not None:
                p.add_run(run.lower()).font.small_caps = True
            else:
                p.add_run(run)

    def process_examples(self, wordDoc, langProps, sentences, translit, translations, additionalInfo,
                         tabular=True, gloss=True, tags=False):
        """
        Add one example to a docx file object.
        """
        if len(sentences) <= 0 or any ('words' not in s for s in sentences):
            return
        if not tabular:
            p = wordDoc.add_paragraph('')
            DocxExampleProcessor.p_no_margins(wordDoc, p)
            for iSent in range(len(sentences)):
                text = sentences[iSent]['text']
                if self.rxLeadingNewline.search(text) is not None:
                    if iSent > 0:
                        # Start new paragraph
                        p = wordDoc.add_paragraph('')
                        DocxExampleProcessor.p_no_margins(wordDoc, p)
                elif iSent > 0:
                    text = ' ' + text.lstrip()
                self.add_pure_text(p, text, translit)
                if self.rxTrailingNewline.search(text) is not None:
                    if iSent < len(sentences) - 1:
                        # Start new paragraph
                        p = wordDoc.add_paragraph('')
                        DocxExampleProcessor.p_no_margins(wordDoc, p)

            for iTier in range(len(translations)):
                p = wordDoc.add_paragraph('‘')
                DocxExampleProcessor.p_no_margins(wordDoc, p)
                for iSent in range(len(translations[iTier])):
                    text = translations[iTier][iSent].strip('‘’')
                    if self.rxLeadingNewline.search(text) is not None:
                        if iSent > 0:
                            # Start new paragraph
                            p = wordDoc.add_paragraph('')
                            DocxExampleProcessor.p_no_margins(wordDoc, p)
                    elif iSent > 0:
                        text = ' ' + text.lstrip()
                    self.add_pure_text(p, text)
                    if self.rxTrailingNewline.search(text) is not None:
                        if iSent < len(translations[iTier]) - 1:
                            # Start new paragraph
                            p = wordDoc.add_paragraph('')
                            DocxExampleProcessor.p_no_margins(wordDoc, p)
                p.text += '’'

        else:
            for iSent in range(len(sentences)):
                curSent = sentences[iSent]
                curTrans = []
                curAddInfo = []
                for iTier in range(len(translations)):
                    if iSent < len(translations[iTier]):
                        curTrans.append(translations[iTier][iSent])
                    else:
                        curTrans.append('')

                for iTier in range(len(additionalInfo)):
                    if iSent < len(additionalInfo[iTier]):
                        curAddInfo.append(additionalInfo[iTier][iSent])
                    else:
                        curAddInfo.append('')

                self.add_glossed_example(wordDoc, langProps, curSent, translit,
                                         curTrans, curAddInfo,
                                         gloss=gloss, tags=tags)
                p = wordDoc.add_paragraph('')
                DocxExampleProcessor.p_no_margins(wordDoc, p)

    def add_pure_text(self, p, text, translit=None):
        """
        Add one example to a paragraph of a docx file object, without glosses.
        """
        # Right now, this is trivial, but maybe we'll add bold/italics
        # and whatnot in the future
        text = self.rxLeadingNewline.sub('', text)
        text = self.rxTrailingNewline.sub('', text)
        if translit is not None:
            p.text += translit(text.rstrip())
        else:
            p.text += text.rstrip()
        p.style = 'Normal'

    def add_glossed_example(self, wordDoc, langProps, s, translit, translations, additionalInfo,
                            gloss=True, tags=False):
        """
        Add one glossed example to a docx file object as a table.
        """
        text = s['text']
        words = []
        glosses = []
        hangingPuncL = ''
        for iToken in range(len(s['words'])):
            w = s['words'][iToken]
            wf = w['wf']
            if w['wtype'] != 'word':
                if self.rxLetters.search(wf) is not None:
                    words.append(hangingPuncL + wf)
                    glosses.append('')
                    hangingPuncL = ''
                else:
                    if w['off_start'] > 0 and text[w['off_start'] - 1] not in (' ', '\t', '\n'):
                        if len(hangingPuncL) > 0 or len(words) <= 0:
                            hangingPuncL += wf
                        else:
                            words[-1] += wf
                    else:
                        hangingPuncL += wf
                continue

            wf = hangingPuncL + wf

            gloss = 'STEM'
            curGlosses = set()
            curParts = set()
            curTrans = set()
            if 'ana' in w:
                for ana in w['ana']:
                    if 'parts' in ana:
                        curParts.add(ana['parts'])
                    if 'gloss' in ana:
                        curGlosses.add(ana['gloss'])
                    if 'trans_en' in ana:
                        curTrans.add(ana['trans_en'])
                    elif 'trans_ru' in ana:
                        curTrans.add(ana['trans_ru'])
            curGlosses = [g for g in sorted(curGlosses, key=lambda x: (x.count('-'), len(x), x))
                          if len(g) > 0]
            curParts = [p for p in sorted(curParts, key=lambda x: (x.count('-'), max(len(p) for p in x.split('-')), x))
                        if len(p) > 0]
            curTrans = [t for t in sorted(curTrans, key=lambda x: (-len(x), x))]
            if len(curGlosses) > 0 and len(curParts) > 0:
                wf = hangingPuncL + curParts[0]
                gloss = curGlosses[0]
                if len(curTrans) > 0:
                    gloss = gloss.replace('STEM', self.rxStemGloss.sub('.', curTrans[0]))
            hangingPuncL = ''

            if ('enclitics_regex' in langProps
                    and langProps['enclitics_regex'].search(wf) is not None
                    and len(words) > 0
                    and len(words[-1]) > 0
                    and self.rxPuncR.search(words[-1][-1]) is None):
                words[-1] += '=' + wf
                glosses[-1] += '=' + gloss
            else:
                words.append(wf)
                glosses.append(gloss)

        nCharsWords = len(''.join(w.strip() for w in words))
        # If a gloss is overly long, something is probably wrong with it,
        # so the user will want to shorten it anyway. Therefore, it would
        # be ok for it to be wrapped
        nCharsGloss = len(''.join(re.sub('^.{30,}',
                                         'X' * 30, g.strip()) for g in glosses))
        nRows = math.ceil(max(nCharsWords // (450 / self.settings.docx_example_font_size),
                              nCharsGloss // (600 / self.settings.docx_gloss_font_size)) + 1)
        nCols = 1 + math.ceil(len(words) / nRows)
        nRows *= 2

        table = wordDoc.add_table(rows=nRows + len(translations) + len(additionalInfo), cols=nCols)
        p = table.cell(0, 0).paragraphs[0]
        p.text = '(xx)'
        DocxExampleProcessor.p_no_margins(wordDoc, p)
        for iRow in range(nRows):
            p = table.cell(iRow, 0).paragraphs[0]
            DocxExampleProcessor.p_no_margins(wordDoc, p)
            # p = table.cell(iRow + 2, 0).paragraphs[0]
            # DocxExampleProcessor.p_no_margins(wordDoc, p)
        for iCell in range(len(words)):
            if iCell >= len(glosses):
                break
            words[iCell] = translit(words[iCell])
            iDoubleRow = iCell // (nCols - 1)
            iCol = iCell - iDoubleRow * (nCols - 1) + 1
            if iCell >= 1:
                for iMergedRow in range(len(translations) + len(additionalInfo)):
                    table.cell(nRows + iMergedRow, 1).merge(table.cell(nRows + iMergedRow, iCol))
            p = table.cell(iDoubleRow * 2, iCol).paragraphs[0]
            p.add_run(words[iCell].strip()).italic = True
            DocxExampleProcessor.p_no_margins(wordDoc, p, 'LanguageExample')
            p = table.cell(iDoubleRow * 2 + 1, iCol).paragraphs[0]
            DocxExampleProcessor.p_no_margins(wordDoc, p, 'Gloss')
            if re.search('^(?:[ /*?!.,()_«»-]*|\\[S[0-9]+\\]:?)$', words[iCell].strip()) is not None:
                continue
            self.smallcaps_glosses(p, glosses[iCell].strip(), langProps)
        for iTrans in range(len(translations)):
            p = table.cell(nRows + iTrans, 1).paragraphs[0]
            p.text = translations[iTrans]
            DocxExampleProcessor.p_no_margins(wordDoc, p)
        for iAddInfo in range(len(additionalInfo)):
            p = table.cell(nRows + len(translations) + iAddInfo, 1).paragraphs[0]
            p.text = additionalInfo[iAddInfo]
            DocxExampleProcessor.p_no_margins(wordDoc, p)
        # self.set_cell_margins(table, 0, 0)
        table.autofit = True

    def extract_translations(self, translations, tabular=True, addQuotes=True):
        """
        Extract and clean translations / additional tiers from lists of source
        dictionaries of the respective sentences.
        """
        # iContext: Contexts = (possibly expanded) search hits
        # iTier: Each context may have multiple translation / additional tiers
        # iSent: Each translation / additional tier is a list of sentences
        for iContext in range(len(translations)):
            for iTier in range(len(translations[iContext])):
                for iSent in range(len(translations[iContext][iTier])):
                    if 'text' in translations[iContext][iTier][iSent]:
                        translations[iContext][iTier][iSent] = translations[iContext][iTier][iSent]['text']
                    else:
                        translations[iContext][iTier][iSent] = ''
                    translations[iContext][iTier][iSent] = translations[iContext][iTier][iSent].strip().replace('\\n', '\n')
                    if tabular:
                        translations[iContext][iTier][iSent] = translations[iContext][iTier][iSent].strip().replace(
                            '\n', ' ')
                        if addQuotes:
                            if not translations[iContext][iTier][iSent].startswith('‘'):
                                translations[iContext][iTier][iSent] = '‘' + translations[iContext][iTier][iSent]
                            if not translations[iContext][iTier][iSent].endswith('’'):
                                translations[iContext][iTier][iSent] += '’'

    def extract_bibref(self, translations):
        references = set()
        for iContext in range(len(translations)):
            for iTier in range(len(translations[iContext])):
                for iSent in range(len(translations[iContext][iTier])):
                    if 'words' not in translations[iContext][iTier][iSent]:
                        continue
                    for w in translations[iContext][iTier][iSent]['words']:
                        if 'bib_ref' in w:
                            references.add(w['bib_ref'])
        return references

    def get_docx(self, lang, contexts, tabular=True, gloss=True, tags=False,
                 translations=None, additionalInfo=None, translit=None,
                 normal_font_face='', normal_font_size=-1,
                 example_font_face='', example_font_size=-1,
                 gloss_font_face='', gloss_font_size=-1):
        """
        Create a docx with the contexts, either as tables or as paragraphs. Each
        context contains one or more sentences in a list.
        """
        if lang not in self.settings.lang_props:
            return None
        langProps = self.settings.lang_props[lang]

        if len(normal_font_face) <= 0:
            normal_font_face = self.settings.docx_normal_font_face
        if normal_font_size <= 0:
            normal_font_face = self.settings.docx_normal_font_size
        if len(example_font_face) <= 0:
            example_font_face = self.settings.docx_example_font_face
        if example_font_size <= 0:
            example_font_size = self.settings.docx_example_font_size
        if len(gloss_font_face) <= 0:
            gloss_font_face = self.settings.docx_gloss_font_face
        if gloss_font_size <= 0:
            gloss_font_size = self.settings.docx_gloss_font_size

        references = set()

        if translations is None:
            translations = []
        references |= self.extract_bibref(translations)
        self.extract_translations(translations, tabular=tabular)
        if additionalInfo is None:
            additionalInfo = []
        references |= self.extract_bibref(additionalInfo)
        self.extract_translations(additionalInfo, tabular=tabular, addQuotes=False)
        if translit is None:
            translit = lambda x: x

        wordDoc = Document()
        exampleStyle = wordDoc.styles.add_style('LanguageExample', WD_STYLE_TYPE.PARAGRAPH)
        glossStyle = wordDoc.styles.add_style('Gloss', WD_STYLE_TYPE.PARAGRAPH)
        normalStyle = wordDoc.styles['Normal']
        normalStyle.font.name = normal_font_face
        normalStyle.font.size = Pt(normal_font_size)
        normalStyle.paragraph_format.space_after = Cm(0)
        exampleStyle.font.name = example_font_face
        exampleStyle.font.size = Pt(example_font_size)
        glossStyle.font.name = gloss_font_face
        glossStyle.font.size = Pt(gloss_font_size)
        for iContext in range(len(contexts)):
            curContext = contexts[iContext]
            for s in curContext:
                if 'words' in s:
                    for w in s['words']:
                        if 'bib_ref' in w:
                            references.add(w['bib_ref'])

            if iContext < len(translations):
                curTrans = translations[iContext]
            else:
                curTrans = []
            if iContext < len(additionalInfo):
                curAddInfo = additionalInfo[iContext]
            else:
                curAddInfo = []
            self.process_examples(wordDoc, langProps, curContext, translit,
                                  curTrans, curAddInfo,
                                  tabular, gloss, tags)
            p = wordDoc.add_paragraph('')
            DocxExampleProcessor.p_no_margins(wordDoc, p)
        if len(references) > 0:
            p = wordDoc.add_paragraph('')
            p.add_run('References').font.bold = True
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Cm(2)
            p.paragraph_format.space_after = Cm(0)
            refFormatted = []
            for ref in references:
                if ref in self.settings.bibref and 'default' in self.settings.bibref[ref]:
                    refFormatted.append(self.settings.bibref[ref]['default'])
                else:
                    refFormatted.append(ref)
            for ref in sorted(refFormatted):
                p = wordDoc.add_paragraph(ref)
                p.paragraph_format.first_line_indent = Cm(1)
                p.paragraph_format.space_before = Cm(0)
                p.paragraph_format.space_after = Cm(0)
                DocxExampleProcessor.p_no_margins(wordDoc, p)
        return wordDoc


if __name__ == '__main__':
    pass

